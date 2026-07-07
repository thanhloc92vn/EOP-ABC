import docx
import os
import sys

def main():
    sys.stdout.reconfigure(encoding='utf-8')
    scratch_dir = os.path.dirname(__file__)
    doc_path = os.path.abspath(os.path.join(scratch_dir, "..", "public", "templates", "bien_ban_hop.docx"))
    out_path = os.path.abspath(os.path.join(scratch_dir, "..", "public", "templates", "bien_ban_hop_template.docx"))
    
    print("Source path:", doc_path)
    print("Output path:", out_path)
    
    if not os.path.exists(doc_path):
        print("Source file not found!")
        return
        
    doc = docx.Document(doc_path)
    
    # 1. Replace paragraphs text
    # Paragraph 4 is Number and Date
    doc.paragraphs[4].text = "Số:  {doc_number}                                        {location_date}"
    
    # Paragraph 6 is Intro block
    doc.paragraphs[6].text = "Hôm nay, vào lúc {start_time} ngày {meeting_date_text}, tại {meeting_location}, Công ty Cổ phần Xây dựng và Lắp máy Trung Nam (TNE&C) đã tiến hành cuộc họp {meeting_title}."
    
    # Paragraph 13 is Closing text
    doc.paragraphs[13].text = "Cuộc họp kết thúc vào lúc {end_time} cùng ngày. Biên bản họp đã được các thành viên thông qua và thực hiện các công việc theo đúng thời gian đã cam kết."
    
    # Paragraph 16 is Nơi nhận values
    doc.paragraphs[16].text = "{distribution}"
    
    # Paragraph 9 is Discussion / Summary
    doc.paragraphs[9].text = "{meeting_summary}\n\nSau khi thảo luận, Chủ trì cuộc họp kết luận và chỉ đạo thực hiện các đầu việc chi tiết như sau:"
    
    # 2. Modify Table 0 (Participants)
    t0 = doc.tables[0]
    t0.rows[0].cells[1].text = "{chair_name}"
    t0.rows[0].cells[2].text = "{chair_role}"
    
    t0.rows[1].cells[1].text = "{sec_name}"
    t0.rows[1].cells[2].text = "{sec_role}"
    
    t0.rows[2].cells[0].text = "Thành viên:"
    t0.rows[2].cells[1].text = "{attendees_text}"
    t0.rows[2].cells[2].text = ""
    
    # Remove rows 3 and 4 from Table 0
    # In python-docx, remove elements from the bottom up
    for i in [4, 3]:
        row = t0.rows[i]
        row._element.getparent().remove(row._element)
        
    # 3. Modify Table 1 (Tasks Table)
    t1 = doc.tables[1]
    
    # In the sample document, Row 1 is a merged row ("MỤC ĐÍCH CUỘC HỌP").
    # Row 2 is the actual unmerged task row.
    t1.rows[2].cells[0].text = "{#tasks}{stt}"
    t1.rows[2].cells[1].text = "{content}"
    t1.rows[2].cells[2].text = "{assignee}"
    t1.rows[2].cells[3].text = "{coop}"
    t1.rows[2].cells[4].text = "{deadline}{/tasks}"
    
    # Remove all rows after row 2
    num_rows = len(t1.rows)
    for i in range(num_rows - 1, 2, -1):
        row = t1.rows[i]
        row._element.getparent().remove(row._element)
        
    # Remove row 1 (the merged "MỤC ĐÍCH CUỘC HỌP" row)
    t1.rows[1]._element.getparent().remove(t1.rows[1]._element)
        
    # 4. Modify Table 2 (Signatures Table)
    t2 = doc.tables[2]
    # Set Thư ký and Chủ trì text
    t2.rows[0].cells[0].text = "THƯ KÝ\n\n\n\n\n{sec_name}"
    t2.rows[0].cells[1].text = "CHỦ TRÌ\n\n\n\n\n{chair_name}"
    
    doc.save(out_path)
    print("Template created successfully!")

if __name__ == "__main__":
    main()
