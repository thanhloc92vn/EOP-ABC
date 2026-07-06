import docx
import os
import sys

def main():
    sys.stdout.reconfigure(encoding='utf-8')
    doc_path = os.path.join(os.path.dirname(__file__), "..", "public", "templates", "bien_ban_hop.docx")
    print("Reading doc:", doc_path)
    
    if not os.path.exists(doc_path):
        print("File does not exist")
        return
        
    doc = docx.Document(doc_path)
    print("Paragraphs:", len(doc.paragraphs))
    print("Tables:", len(doc.tables))
    
    print("\n--- Paragraphs ---")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"[{i}]: {p.text.strip()}")
            
    print("\n--- Tables ---")
    for i, t in enumerate(doc.tables):
        print(f"Table {i}: rows={len(t.rows)}, cols={len(t.columns)}")
        for r_idx, r in enumerate(t.rows[:5]): # first 5 rows
            row_text = [cell.text.strip().replace('\n', ' ') for cell in r.cells]
            print(f"  Row {r_idx}: {row_text}")

if __name__ == "__main__":
    main()
