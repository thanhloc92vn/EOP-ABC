import docx

doc = docx.Document("dashboard/scratch/test_out.docx")

with open("dashboard/inspect_out_test.txt", "w", encoding="utf-8") as f:
    f.write("=== GENERATED DOCX PARAGRAPH 0 DETAILED RUNS ===\n")
    p = doc.paragraphs[0]
    f.write(f"Paragraph text: '{p.text}'\n")
    for idx, run in enumerate(p.runs):
        f.write(f"  Run {idx} text: '{run.text}'\n")
        f.write(f"    - has drawing: {'drawing' in run._r.xml}\n")
        f.write(f"    - has pic: {'pic' in run._r.xml}\n")
        
    f.write("\n=== ALL PARAGRAPHS ===\n")
    for i, para in enumerate(doc.paragraphs):
        f.write(f"Paragraph {i}: '{para.text}'\n")
        
print("Inspection written to dashboard/inspect_out_test.txt")
