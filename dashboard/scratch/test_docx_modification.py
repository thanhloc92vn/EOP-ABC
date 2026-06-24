import docx
import datetime

doc = docx.Document("dashboard/public/templates/bao_cao_tuyen_dung_tuan.docx")

# Prepare dates
start_date = datetime.date(2026, 6, 1)
end_date = datetime.date(2026, 6, 11)

formatted_start = start_date.strftime("%d/%m")
formatted_end = end_date.strftime("%d/%m/%Y")

date_range_str = f" (Từ ngày {formatted_start} đến ngày {formatted_end})"

# Modify title
title_modified = False
for p in doc.paragraphs:
    if "BÁO CÁO TUYỂN DỤNG" in p.text or "BÁO CÁO TUYỂN DỤNG THÁNG" in p.text:
        for run in p.runs:
            if "BÁO CÁO TUYỂN DỤNG" in run.text or "BÁO CÁO TUYỂN DỤNG THÁNG" in run.text:
                run.text = f"BÁO CÁO TUYỂN DỤNG TUẦN\n(Từ ngày {formatted_start} đến ngày {formatted_end})"
                title_modified = True
                print("Modified title run text successfully!")
                break

# Modify bottom date
date_modified = False
day_str = str(end_date.day).zfill(2)
month_str = str(end_date.month).zfill(2)
year_str = str(end_date.year)

for p in doc.paragraphs:
    if "ngày" in p.text.lower() and "tháng" in p.text.lower() and "năm" in p.text.lower() and len(p.text) < 50:
        # Instead of replacing p.text which clears all runs, let's modify the runs or p.text
        # Let's see if we can just set p.text and preserve/apply style
        p.text = f"Ngày {day_str} tháng {month_str} năm {year_str}"
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            run.italic = True
            run.font.name = "Arial"
        date_modified = True
        print("Modified bottom date successfully!")
        break

doc.save("dashboard/scratch/test_out.docx")
print("Saved to dashboard/scratch/test_out.docx")

# Inspect the saved docx to verify runs in Paragraph 0
doc2 = docx.Document("dashboard/scratch/test_out.docx")
print("\nAfter modification Paragraph 0 runs:")
for idx, run in enumerate(doc2.paragraphs[0].runs):
    print(f"  Run {idx}: text='{run.text}'")
    if "drawing" in run._r.xml:
        print("    - Still contains drawing/picture XML")
