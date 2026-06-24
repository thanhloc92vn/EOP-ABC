import docx
import os
import sys

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

root_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
out_path = os.path.join(root_dir, "scratch", "test_report_out.docx")

print("Output docx exists:", os.path.exists(out_path))
if os.path.exists(out_path):
    doc = docx.Document(out_path)
    print("Number of tables:", len(doc.tables))
    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- TABLE {t_idx} ---")
        for r_idx, row in enumerate(table.rows):
            row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            print(f"Row {r_idx}: {row_text}")
